#!/usr/bin/env python3
"""Python/Automation skills tracker - weekly lessons on Tue/Thu."""
import os
import json
import sys
from datetime import datetime

STATE_FILE = os.path.expanduser("~/.hermes/cron/python_skills_state.json")

TOPICS = [
    {
        "day": 1,
        "title": "🐍 文件批量处理 — os + pathlib",
        "content": (
            "**场景：** 每天处理一堆文件，手动改太慢\n\n"
            "**核心模块：**\n"
            "```python\n"
            "import os\n"
            "from pathlib import Path\n"
            "\n"
            "# 遍历目录下所有.txt文件\n"
            "for f in Path('./logs').glob('*.txt'):\n"
            "    print(f.name)  # 文件名\n"
            "    print(f.suffix)  # 扩展名\n"
            "    print(f.stat().st_size)  # 文件大小\n"
            "\n"
            "# 批量重命名：添加日期前缀\n"
            "from datetime import date\n"
            "today = date.today().strftime('%Y%m%d')\n"
            "for f in Path('./reports').glob('*.xlsx'):\n"
            "    f.rename(f.parent / f'{today}_{f.name}')\n"
            "```\n\n"
            "**练手任务：** 写一个脚本，把Downloads目录下所有.png文件按月份移动到 ./images/2026-05/ 这种文件夹里"
        ),
        "practice": (
            "**🛠 今日练习：批量整理文件**\n\n"
            "需求：桌面 ~/Desktop 下有一堆截图文件（screenshot_001.png, screenshot_002.png ...），写一个脚本：\n"
            "1. 找到所有 .png 文件\n"
            "2. 按文件修改日期放入 ~/Desktop/screenshots/2026-05/ 等月份文件夹\n"
            "3. 打印每个文件移动前后的路径\n\n"
            "**提示：** `os.path.getmtime()` 或 `Path.stat().st_mtime` 获取修改时间\n"
            "**参考：**\n"
            "```python\n"
            "import time\n"
            "from pathlib import Path\n"
            "from datetime import datetime\n"
            "\n"
            "src = Path.home() / 'Desktop'\n"
            "for f in src.glob('*.png'):\n"
            "    mtime = datetime.fromtimestamp(f.stat().st_mtime)\n"
            "    month_dir = src / 'screenshots' / mtime.strftime('%Y-%m')\n"
            "    month_dir.mkdir(parents=True, exist_ok=True)\n"
            "    f.rename(month_dir / f.name)\n"
            "    print(f'移动: {f} -> {month_dir / f.name}')\n"
            "```\n\n"
            "💡 下周推送到你手机上打卡完成情况～"
        )
    },
    {
        "day": 2,
        "title": "🐍 正则表达式 — 信息提取利器",
        "content": (
            "**场景：** 从日志里提取IP、抓网页里所有链接、格式化手机号\n\n"
            "**核心模块：** `re`\n\n"
            "```python\n"
            "import re\n"
            "\n"
            "# 从log里提取所有IP地址\n"
            "log = \"ERROR 192.168.1.100 connection timeout, from 10.0.0.5\"\n"
            "ips = re.findall(r'\\d+\\.\\d+\\.\\d+\\.\\d+', log)\n"
            "print(ips)  # ['192.168.1.100', '10.0.0.5']\n"
            "\n"
            "# 提取AT指令返回值\n"
            "resp = \"+CSQ: 28,99\"\n"
            "m = re.search(r'\\+CSQ: (\\d+),(\\d+)', resp)\n"
            "if m:\n"
            "    rssi, ber = m.group(1), m.group(2)\n"
            "    print(f'信号强度: {rssi}, 误码率: {ber}')\n"
            "\n"
            "# 手机号脱敏\n"
            "phone = \"13812345678\"\n"
            "masked = re.sub(r'(\\d{3})\\d{4}(\\d{4})', r'\\1****\\2', phone)\n"
            "print(masked)  # 138****5678\n"
            "```\n\n"
            "**常用正则速记：**\n"
            "- `\\d` 数字、`\\w` 字母数字下划线、`\\s` 空白\n"
            "- `+` 一次或多次、`*` 零次或多次、`?` 零次或一次\n"
            "- `.*?` 非贪婪匹配（很重要！不加?会匹配到最后一个）"
        ),
        "practice": (
            "**🛠 今日练习：日志IP提取**\n\n"
            "需求：公司服务器log如下，提取所有IP并去重：\n\n"
            "```\n"
            "May 20 10:23:45 server sshd[1234]: Failed password for root from 192.168.1.100 port 22\n"
            "May 20 10:24:01 server sshd[1235]: Accepted password for admin from 10.0.0.88 port 22\n"
            "May 20 10:25:12 server sshd[1236]: Failed password for root from 192.168.1.100 port 22\n"
            "May 20 10:26:30 server sshd[1237]: Failed password for root from 203.0.113.50 port 22\n"
            "```\n\n"
            "```python\n"
            "import re\n"
            "logs = \"\"\"...上面那段...\"\"\"\n"
            "ips = set(re.findall(r'\\d+\\.\\d+\\.\\d+\\.\\d+', logs))\n"
            "for ip in sorted(ips):\n"
            "    print(ip)\n"
            "```\n\n"
            "💡 扩展：统计每个IP出现的次数 → 用 collections.Counter"
        )
    },
    {
        "day": 3,
        "title": "🐍 Excel自动化 — openpyxl",
        "content": (
            "**场景：** 每周整理报表、合并多个Excel、批量修改格式\n\n"
            "```bash\n"
            "pip install openpyxl\n"
            "```\n\n"
            "```python\n"
            "from openpyxl import load_workbook, Workbook\n"
            "\n"
            "# 读取已有文件\n"
            "wb = load_workbook('report.xlsx')\n"
            "ws = wb.active  # 当前工作表\n"
            "\n"
            "# 遍历每一行\n"
            "for row in ws.iter_rows(min_row=2, values_only=True):\n"
            "    name, score = row[0], row[1]\n"
            "    print(f'{name}: {score}')\n"
            "\n"
            "# 新建Excel并写入\n"
            "wb2 = Workbook()\n"
            "ws2 = wb2.active\n"
            "ws2.title = '销售数据'\n"
            "ws2['A1'] = '产品'\n"
            "ws2['B1'] = '销量'\n"
            "ws2.append(['模组A', 1200])\n"
            "ws2.append(['模组B', 850])\n"
            "wb2.save('output.xlsx')\n"
            "```\n\n"
            "**实用技巧：**\n"
            "```python\n"
            "# 合并多个Excel到一个文件\n"
            "from pathlib import Path\n"
            "\n"
            "all_data = []\n"
            "for f in Path('.').glob('sales_*.xlsx'):\n"
            "    wb = load_workbook(f)\n"
            "    ws = wb.active\n"
            "    for row in ws.iter_rows(min_row=2, values_only=True):\n"
            "        all_data.append(row)\n",
            "```"
        ),
        "practice": (
            "**🛠 今日练习：合并月报**\n\n"
            "需求：目录下有 weekly_report_1.xlsx 到 weekly_report_4.xlsx 四个文件，\n"
            "每个都有表头（日期、产品、销量、金额），合并成一个汇总文件。\n\n"
            "```python\n"
            "from openpyxl import load_workbook, Workbook\n"
            "from pathlib import Path\n"
            "\n"
            "wb_out = Workbook()\n"
            "ws_out = wb_out.active\n"
            "ws_out.append(['日期', '产品', '销量', '金额'])  # 写表头\n"
            "\n"
            "for f in sorted(Path('.').glob('weekly_report_*.xlsx')):\n"
            "    wb = load_workbook(f)\n"
            "    ws = wb.active\n"
            "    for row in ws.iter_rows(min_row=2, values_only=True):\n"
            "        ws_out.append(row)\n"
            "\n"
            "wb_out.save('monthly_summary.xlsx')\n"
            "print(f'合并完成！共 {ws_out.max_row - 1} 条数据')\n"
            "```\n\n"
            "💡 实际工作中这个脚本每月省你半小时手动复制粘贴"
        )
    },
    {
        "day": 4,
        "title": "🐍 网络请求 — requests 入门",
        "content": (
            "**场景：** 自动抓网页数据、调用API接口、检测网站是否正常\n\n"
            "```bash\n"
            "pip install requests beautifulsoup4\n"
            "```\n\n"
            "```python\n"
            "import requests\n"
            "\n"
            "# GET请求 - 抓网页\n"
            "resp = requests.get('https://api.github.com')\n"
            "print(resp.status_code)  # 200\n"
            "print(resp.json())  # 如果是JSON自动解析\n"
            "\n"
            "# 带参数\n"
            "params = {'wd': '移远通信', 'rn': 10}\n"
            "resp = requests.get('https://www.baidu.com/s', params=params)\n"
            "\n"
            "# POST请求\n"
            "data = {'username': 'admin', 'password': '123456'}\n"
            "resp = requests.post('https://httpbin.org/post', data=data)\n"
            "\n"
            "# 设置请求头（模拟浏览器）\n"
            "headers = {\n"
            "    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'\n"
            "}\n"
            "resp = requests.get(url, headers=headers)\n"
            "```\n\n"
            "**错误处理：**\n"
            "```python\n"
            "try:\n"
            "    resp = requests.get(url, timeout=5)\n"
            "    resp.raise_for_status()  # 非200状态码会抛异常\n"
            "except requests.exceptions.ConnectionError:\n"
            "    print('网络不通')\n"
            "except requests.exceptions.Timeout:\n"
            "    print('请求超时')\n"
            "except requests.exceptions.HTTPError as e:\n"
            "    print(f'HTTP错误: {e}')\n"
            "```"
        ),
        "practice": (
            "**🛠 今日练习：股票查询小工具**\n\n"
            "需求：用新浪财经API查某只A股实时价格\n\n"
            "```python\n"
            "import requests\n"
            "\n"
            "# 上证指数代码: sh000001\n"
            "# 贵州茅台代码: sh600519\n"
            "code = 'sh600519'\n"
            "url = f'https://hq.sinajs.cn/list={code}'\n"
            "\n"
            "headers = {'Referer': 'https://finance.sina.com.cn'}\n"
            "resp = requests.get(url, headers=headers)\n"
            "resp.encoding = 'gbk'\n"
            "\n"
            "# 返回格式: var hq_str_sh600519=\"茅台,1450.00,1455.00,...\"\n"
            "data = resp.text.split('=')[1].strip().strip('\"').split(',')\n"
            "name = data[0]\n"
            "price = data[1]\n"
            "change = data[3]\n"
            "print(f'{name}: ¥{price} (涨跌: {change})')\n"
            "```\n\n"
            "💡 试试改成查询多只股票，一次性输出你的持仓"
        )
    },
    {
        "day": 5,
        "title": "🐍 PDF/DOCX 处理 — 办公必备",
        "content": (
            "**场景：** 批量提取PDF文字、生成报告、合并文档\n\n"
            "```bash\n"
            "pip install python-docx PyPDF2\n"
            "```\n\n"
            "**读取PDF：**\n"
            "```python\n"
            "from PyPDF2 import PdfReader\n"
            "\n"
            "reader = PdfReader('report.pdf')\n"
            "print(f'共{len(reader.pages)}页')\n"
            "for page in reader.pages:\n"
            "    text = page.extract_text()\n"
            "    print(text[:200])  # 前200字\n"
            "```\n\n"
            "**生成Word报告：**\n"
            "```python\n"
            "from docx import Document\n"
            "\n"
            "doc = Document()\n"
            "doc.add_heading('移远技术支持日报', level=1)\n"
            "doc.add_paragraph('今日处理客户问题3个，已全部闭环。')\n"
            "\n"
            "# 加表格\n"
            "table = doc.add_table(rows=3, cols=3)\n"
            "table.style = 'Light Grid'\n"
            "for i, text in enumerate(['客户', '问题', '状态']):\n"
            "    table.rows[0].cells[i].text = text\n"
            "table.rows[1].cells[0].text = '客户A'\n"
            "table.rows[1].cells[1].text = '模组无法注册'\n"
            "table.rows[1].cells[2].text = '已解决'\n"
            "doc.save('日报.docx')\n"
            "```"
        ),
        "practice": (
            "**🛠 今日练习：自动生成日报**\n\n"
            "需求：每天重复写日报很烦，写个脚本自动生成\n\n"
            "```python\n"
            "from docx import Document\n"
            "from datetime import date\n"
            "\n"
            "today = date.today()\n"
            "doc = Document()\n"
            "doc.add_heading(f'技术支持日报 - {today}', level=1)\n"
            "\n"
            "# 今天处理的问题\n"
            "issues = [\n"
            "    ('客户A', '5G模组无法搜网', '已解决', 'SIM卡接触不良'),\n"
            "    ('客户B', 'AT指令无响应', '处理中', '等待客户提供log'),\n"
            "]\n"
            "\n"
            "table = doc.add_table(rows=1+len(issues), cols=4)\n"
            "table.style = 'Light Grid'\n"
            "headers = ['客户', '问题', '状态', '备注']\n"
            "for i, h in enumerate(headers):\n"
            "    table.rows[0].cells[i].text = h\n"
            "for r, issue in enumerate(issues):\n"
            "    for c, val in enumerate(issue):\n"
            "        table.rows[r+1].cells[c].text = val\n"
            "\n"
            "doc.save(f'日报_{today}.docx')\n"
            "print('日报已生成！')\n"
            "```\n\n"
            "💡 把这个脚本放cron，每天下班前自动发给你检查"
        )
    },
    {
        "day": 6,
        "title": "🐍 数据快速分析 — pandas入门",
        "content": (
            "**场景：** 分析Excel/CSV数据、出统计报表、做图表\n\n"
            "```bash\n"
            "pip install pandas openpyxl\n"
            "```\n\n"
            "```python\n"
            "import pandas as pd\n"
            "\n"
            "# 读取CSV/Excel\n"
            "df = pd.read_csv('sales.csv')\n"
            "# df = pd.read_excel('sales.xlsx')\n"
            "\n"
            "# 看一眼数据\n"
            "print(df.head())  # 前5行\n"
            "print(df.info())  # 列名、类型、非空计数\n"
            "print(df.describe())  # 数值列的统计摘要\n"
            "\n"
            "# 筛选\n"
            "high_sales = df[df['销量'] > 1000]\n"
            "product_a = df[df['产品'].str.contains('模组A')]\n"
            "\n"
            "# 分组统计\n"
            "summary = df.groupby('产品')['销量'].sum()\n"
            "print(summary)\n"
            "\n"
            "# 排序\n"
            "top5 = df.sort_values('销量', ascending=False).head(5)\n"
            "\n"
            "# 导出\n"
            "summary.to_excel('统计结果.xlsx')\n"
            "```\n\n"
            "**面试加分技巧：** pandas处理10万行数据也飞快，相比Excel手动操作效率天差地别"
        ),
        "practice": (
            "**🛠 今日练习：销售数据分析**\n\n"
            "假设有以下CSV数据（sales.csv）：\n"
            "```\n"
            "日期,产品,销量,金额\n"
            "2026-05-01,模组A,1200,48000\n"
            "2026-05-01,模组B,800,32000\n"
            "2026-05-02,模组A,1500,60000\n"
            "2026-05-02,模组B,600,24000\n"
            "2026-05-03,模组A,900,36000\n"
            "2026-05-03,模组B,1100,44000\n"
            "```\n\n"
            "```python\n"
            "import pandas as pd\n"
            "\n"
            "df = pd.read_csv('sales.csv')\n"
            "\n"
            "# 1. 哪个产品总销量最高？\n"
            "print(df.groupby('产品')['销量'].sum())\n"
            "\n"
            "# 2. 哪天销售额最高？\n"
            "df['日期'] = pd.to_datetime(df['日期'])\n"
            "daily = df.groupby('日期')['金额'].sum()\n"
            "print(f'最高销售日: {daily.idxmax().date()} ¥{daily.max()}')\n"
            "\n"
            "# 3. 平均单价\n"
            "df['单价'] = df['金额'] / df['销量']\n"
            "print(df.groupby('产品')['单价'].mean())\n"
            "```"
        )
    },
    {
        "day": 7,
        "title": "🐍 定时任务 — 用Python自动干活",
        "content": (
            "**场景：** 每天自动发邮件、定时备份、定期检查\n\n"
            "**方案1：Python schedule库**\n"
            "```bash\n"
            "pip install schedule\n"
            "```\n"
            "```python\n"
            "import schedule\n"
            "import time\n"
            "\n"
            "def job():\n"
            "    print('正在执行定时任务...')\n"
            "\n"
            "schedule.every().day.at('09:00').do(job)\n"
            "schedule.every(30).minutes.do(job)\n"
            "schedule.every().monday.at('10:00').do(job)\n"
            "\n"
            "while True:\n"
            "    schedule.run_pending()\n"
            "    time.sleep(1)\n"
            "```\n\n"
            "**方案2：直接用cron（推荐）**\n"
            "```bash\n"
            "# 每天9点执行Python脚本\n"
            "0 9 * * * cd /path && python3 script.py\n",
            "```\n\n"
            "**实用案例：每天自动备份**\n"
            "```python\n"
            "import shutil\n"
            "from pathlib import Path\n"
            "from datetime import datetime\n"
            "\n"
            "today = datetime.now().strftime('%Y%m%d')\n"
            "src = Path.home() / 'Documents'\n"
            "dst = Path.home() / 'backups' / f'docs_{today}'\n"
            "shutil.copytree(src, dst)\n"
            "print(f'备份完成: {dst}')\n"
            "```"
        ),
        "practice": (
            "**🛠 今日综合练习：自动备份脚本**\n\n"
            "写一个每天早上9点运行的脚本：\n"
            "1. 压缩 ~/Documents/工作日报/ 目录为ZIP\n"
            "2. 保存到 ~/backups/日报备份_YYYYMMDD.zip\n"
            "3. 只保留最近7天的备份（删除旧的）\n\n"
            "```python\n"
            "from pathlib import Path\n"
            "import shutil\n"
            "from datetime import datetime, timedelta\n"
            "\n"
            "today = datetime.now().strftime('%Y%m%d')\n"
            "src = Path.home() / 'Documents' / '工作日报'\n"
            "backup_dir = Path.home() / 'backups'\n"
            "backup_dir.mkdir(exist_ok=True)\n"
            "\n"
            "# 压缩备份\n"
            "zip_path = backup_dir / f'日报备份_{today}.zip'\n"
            "shutil.make_archive(str(zip_path.with_suffix('')), 'zip', src)\n"
            "print(f'备份: {zip_path}')\n"
            "\n"
            "# 删除7天前的\n"
            "for f in backup_dir.glob('日报备份_*.zip'):\n"
            "    date_str = f.stem.split('_')[1]\n"
            "    file_date = datetime.strptime(date_str, '%Y%m%d')\n"
            "    if file_date < datetime.now() - timedelta(days=7):\n"
            "        f.unlink()\n"
            "        print(f'删除旧备份: {f}')\n"
            "```"
        )
    },
    {
        "day": 8,
        "title": "🐍 综合实战 — 爬取+分析+报表一条龙",
        "content": (
            "**场景：** 把前面学的串起来，做一个完整项目\n\n"
            "**实战项目：竞品价格监控**\n"
            "```python\n"
            "import requests\n"
            "import pandas as pd\n"
            "from datetime import datetime\n"
            "from pathlib import Path\n"
            "\n"
            "# 1. 抓数据\n"
            "products = ['模组A', '模组B', '模组C']\n"
            "prices = []\n"
            "for p in products:\n",
            "    # 模拟查价格（实际换成真实的API）\n",
            "    prices.append({'产品': p, '价格': 120 + len(prices) * 10})\n",
            "\n"
            "# 2. 存到Excel\n"
            "df = pd.DataFrame(prices)\n",
            "df['日期'] = datetime.now().strftime('%Y-%m-%d')\n",
            "today_str = datetime.now().strftime('%Y%m%d')\n",
            "df.to_excel(f'price_{today_str}.xlsx', index=False)\n",
            "\n",
            "# 3. 和历史数据合并\n",
            "hist_file = 'price_history.xlsx'\n",
            "if Path(hist_file).exists():\n",
            "    hist = pd.read_excel(hist_file)\n",
            "    df = pd.concat([hist, df])\n",
            "df.to_excel(hist_file, index=False)\n",
            "\n",
            "print(f'价格监控完成，共 {len(df)} 条记录')\n",
            "```\n\n"
            "**💡 核心竞争力：** 别人手动查价格填表格，你一个脚本自动搞定。这就是用Python提升价值的体现"
        ),
        "practice": (
            "**🛠 今日实战：构建你的工具库**\n\n"
            "把你学过的脚本整理成一个 `toolkit.py`：\n"
            "```python\n"
            "#!/usr/bin/env python3\n"
            "\"\"\"我的办公自动化工具箱\"\"\"\n"
            "\n"
            "def batch_rename(path, prefix, ext):\n",
            "    \"批量重命名文件\"\n",
            "    ...\n\n"
            "def extract_ips(log_text):\n",
            "    \"提取日志中所有IP\"\n",
            "    ...\n\n"
            "def merge_excel(pattern, output):\n",
            "    \"合并多个Excel\"\n"
            "    ...\n\n"
            "def gen_daily_report(issues, output):\n",
            "    \"生成日报Word\"\n",
            "    ...\n\n"
            "def analyze_sales(csv_file):\n",
            "    \"销售数据分析\"\n",
            "    ...\n\n"
            "if __name__ == '__main__':\n",
            "    # 命令行入口\n",
            "    import sys\n",
            "    if len(sys.argv) > 1:\n",
            "        cmd = sys.argv[1]\n",
            "        if cmd == 'rename':\n",
            "            batch_rename(...)\n",
            "```\n\n"
            "💡 以后重复工作直接 `python toolkit.py xxx` 一键搞定\n",
            "面试时也可以展示自己的自动化项目经验"
        )
    }
]

def load_state():
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, 'r') as f:
            return json.load(f)
    return {"current_day": 0, "started_at": None, "last_morning_date": None}

def save_state(state):
    os.makedirs(os.path.dirname(STATE_FILE), exist_ok=True)
    with open(STATE_FILE, 'w') as f:
        json.dump(state, f, indent=2)

def get_morning_topic():
    state = load_state()
    today = datetime.now().strftime("%Y-%m-%d")
    
    if state.get("last_morning_date") == today:
        return TOPICS[state["current_day"] - 1] if state["current_day"] > 0 else None
    
    state["current_day"] = state.get("current_day", 0) + 1
    state["last_morning_date"] = today
    if state.get("started_at") is None:
        state["started_at"] = today
    
    if state["current_day"] > len(TOPICS):
        state["current_day"] = 1
    
    topic = TOPICS[state["current_day"] - 1]
    save_state(state)
    return topic

def get_evening_topic():
    state = load_state()
    if state.get("current_day", 0) == 0:
        return None
    return TOPICS[state["current_day"] - 1]

if __name__ == "__main__":
    mode = sys.argv[1] if len(sys.argv) > 1 else "morning"
    
    if mode == "morning":
        topic = get_morning_topic()
    else:
        topic = get_evening_topic()
    
    if topic:
        total = len(TOPICS)
        print(f"LESSON_DAY={topic['day']}")
        print(f"LESSON_TOTAL={total}")
        print(f"LESSON_TITLE={topic['title']}")
        if mode == "morning":
            print(f"LESSON_CONTENT={topic['content']}")
        else:
            print(f"LESSON_CONTENT={topic['practice']}")
